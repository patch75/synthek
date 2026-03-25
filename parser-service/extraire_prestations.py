"""
Synthek — Extraction des prestations depuis une notice descriptive
Utilise Sonnet 4.6 pour extraire les prestations par financement.
"""
import os
import re
import json
import io
import base64
import anthropic
import pdfplumber
from docx import Document as DocxDocument

MODELE_SONNET = "claude-sonnet-4-6"

# Valeurs autorisées par champ — validation Python après extraction
VALEURS_AUTORISEES = {
    "chauf_distribution": ["individuel", "collectif", "mixte"],
    "chauf_production": [
        "pac_air_eau", "pac_eau_eau", "chaudiere_gaz_individuelle",
        "chaudiere_gaz_collective", "chaudiere_biomasse", "rcu", "effet_joule"
    ],
    "chauf_emetteurs": [
        "plt", "radiateurs_eau", "ventiloconvecteurs", "plafond_rayonnant",
        "convecteurs_electriques", "seche_serviettes_elec"
    ],
    "chauf_regulation": [
        "robinets_thermostatiques", "gtb_gtc", "thermostat_ambiance", "programmation_logement"
    ],
    "ecs_production": [
        "pac_thermo_individ", "pac_thermo_collectif", "chauffe_eau_elec",
        "cesi", "scsc", "chaudiere_gaz_individuelle", "rcu", "ballon_elec_collectif"
    ],
    "ecs_distribution": [
        "individuelle_sans_boucle", "boucle_collective", "bouclage_anti_legionellose"
    ],
    "vmc_type": [
        "sf_a", "sf_hygro_a", "sf_hygro_b", "df_collectif", "df_individuel"
    ],
    "san_wc": ["sol", "suspendu"],
    "san_vasque": ["meuble", "poser", "suspendu", "attente_seulement"],
    "san_douche": ["receveur_extra_plat_paroi", "receveur_standard", "italienne"],
    "san_baignoire": ["encastree_pare_baignoire", "selon_plan", "aucune"],
    "san_robinetterie": ["mecanique", "thermostatique", "electronique"],
    "enr_type": ["solaire_th", "pv", "geothermie", "aucune"],
}

# Champs qui acceptent plusieurs valeurs (JSON array string)
CHAMPS_ARRAY = {"chauf_emetteurs", "chauf_regulation"}

PROMPT_SYSTEME = """Tu es un expert en notices descriptives de logements collectifs (France).
Tu extrais les prestations techniques d'une notice selon un financement donné.

RÈGLES CRITIQUES :
1. Lire en priorité les PARTIES PRIVATIVES — ignorer parties communes pour les prestations logement.
   Piège notices : le générateur peut être décrit vaguement p.8 (parties communes) puis précisément
   p.12 (parties privatives) → prioriser la description la plus précise et la plus détaillée.
2. Distinguer "fourni" vs "en attente" → attente seulement = null, ne jamais extrapoler.
3. Valeur conditionnelle ("selon plan", "selon configuration") → fiabilite: "a_confirmer".
4. Contradiction entre deux pages → prioriser la page la plus précise.
5. Si une information n'est pas dans le document → null (ne pas inventer).

VALEURS ACCEPTÉES PAR CHAMP (utiliser EXACTEMENT ces valeurs) :
- chauf_distribution: "individuel" | "collectif" | "mixte"
- chauf_production: "pac_air_eau" | "pac_eau_eau" | "chaudiere_gaz_individuelle" | "chaudiere_gaz_collective" | "chaudiere_biomasse" | "rcu" | "effet_joule"
- chauf_emetteurs: ["plt","radiateurs_eau","ventiloconvecteurs","plafond_rayonnant","convecteurs_electriques","seche_serviettes_elec"] — TABLEAU (plusieurs possibles)
- chauf_regulation: ["robinets_thermostatiques","gtb_gtc","thermostat_ambiance","programmation_logement"] — TABLEAU
- ecs_production: "pac_thermo_individ" | "pac_thermo_collectif" | "chauffe_eau_elec" | "cesi" | "scsc" | "chaudiere_gaz_individuelle" | "rcu" | "ballon_elec_collectif"
- ecs_distribution: "individuelle_sans_boucle" | "boucle_collective" | "bouclage_anti_legionellose"
- vmc_type: "sf_a" | "sf_hygro_a" | "sf_hygro_b" | "df_collectif" | "df_individuel"
- san_wc: "sol" | "suspendu"
- san_vasque: "meuble" | "poser" | "suspendu" | "attente_seulement"
- san_douche: "receveur_extra_plat_paroi" | "receveur_standard" | "italienne"
- san_baignoire: "encastree_pare_baignoire" | "selon_plan" | "aucune"
- san_robinetterie: "mecanique" | "thermostatique" | "electronique"
- enr_type: "solaire_th" | "pv" | "geothermie" | "aucune"

RÉPONDRE UNIQUEMENT avec ce JSON (pas de markdown, pas d'explication) :
{
  "chauf_distribution": "...",
  "chauf_production": "...",
  "chauf_emetteurs": ["..."],
  "chauf_regulation": ["..."],
  "ecs_production": "...",
  "ecs_distribution": "...",
  "vmc_type": "...",
  "san_wc": "...",
  "san_vasque": "...",
  "san_douche": "...",
  "san_baignoire": "...",
  "san_robinetterie": "...",
  "enr_type": "...",
  "fiabilite": "haute" | "moyenne" | "a_confirmer",
  "noteComplementaire": "..." ou null
}"""


def _extraire_texte_pdf(file_bytes: bytes) -> str:
    texte = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            t = page.extract_text()
            if t:
                texte.append(f"--- Page {i + 1} ---\n{t}")
    return "\n\n".join(texte)


def _extraire_texte_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    lignes = []
    for para in doc.paragraphs:
        if para.text.strip():
            lignes.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lignes.append(" | ".join(cells))
    return "\n".join(lignes)


def _valider_et_nettoyer(data: dict) -> dict:
    """Valide les valeurs extraites et nettoie les invalides (→ None)."""
    for champ, valeurs_ok in VALEURS_AUTORISEES.items():
        val = data.get(champ)
        if val is None:
            continue
        if champ in CHAMPS_ARRAY:
            if isinstance(val, list):
                valides = [v for v in val if v in valeurs_ok]
                data[champ] = json.dumps(valides) if valides else None
            elif isinstance(val, str):
                # Sonnet a renvoyé une string au lieu d'un tableau
                valides = [v for v in [val] if v in valeurs_ok]
                data[champ] = json.dumps(valides) if valides else None
            else:
                data[champ] = None
        else:
            if val not in valeurs_ok:
                data[champ] = None
    return data


def extraire_prestations(file_bytes: bytes, nom_fichier: str, financement: str) -> dict:
    """
    Point d'entrée principal.
    Retourne un dict avec les champs prestations + fiabilite + source.
    """
    nom_lower = nom_fichier.lower()
    if nom_lower.endswith('.pdf'):
        texte = _extraire_texte_pdf(file_bytes)
    elif nom_lower.endswith('.docx'):
        texte = _extraire_texte_docx(file_bytes)
    else:
        raise ValueError(f"Format non supporté : {nom_fichier}. Utiliser PDF ou DOCX.")

    if not texte.strip():
        raise ValueError("Aucun texte extrait du document.")

    # Tronquer si trop long (Sonnet 4.6 : 200k context, mais on limite pour le coût)
    LIMITE = 40000
    if len(texte) > LIMITE:
        texte = texte[:LIMITE] + "\n\n[... document tronqué ...]"

    FINANCEMENT_LABEL = {
        "social": "Social (LLI/LLS/LLTS/PLS)",
        "brs": "BRS (Bail Réel Solidaire)",
        "acces_std": "Accession standard",
        "premium": "Accession premium / Attique",
    }
    label_financement = FINANCEMENT_LABEL.get(financement, financement)

    client = anthropic.Anthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
    message = client.messages.create(
        model=MODELE_SONNET,
        max_tokens=1500,
        system=PROMPT_SYSTEME,
        messages=[{
            "role": "user",
            "content": (
                f"Notice descriptive : {nom_fichier}\n"
                f"Financement à extraire : {label_financement}\n\n"
                f"{texte}"
            )
        }]
    )

    raw = message.content[0].text.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        idx = raw.find('{')
        if idx >= 0:
            data = json.loads(raw[idx:])
        else:
            raise ValueError(f"Sonnet n'a pas retourné un JSON valide : {raw[:300]}")

    data = _valider_et_nettoyer(data)
    data["financement"] = financement
    data["source"] = "notice"

    return data
