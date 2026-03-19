"""
Synthek Parser Service — microservice local (port 5001)
Parsing amélioré de fichiers DPGF/CCTP/PDF avec reconstruction hiérarchie parent/enfant.
"""
from flask import Flask, request, jsonify
import openpyxl
from docx import Document
import pdfplumber
import base64
import io
import traceback
from comparaison_cctp_dpgf import extraire_cctp, extraire_dpgf, detecter_alertes, extraire_programme
from equivalences_fluides import est_ligne_exclue, sont_equivalents

app = Flask(__name__)


def parse_xlsx(file_bytes):
    """
    Parse Excel DPGF avec reconstruction de la hiérarchie parent/enfant.
    Les lignes-sections (ex: "Vanne d'arrêt générale repérée") sont préfixées
    aux lignes-détail (ex: "DN 40 : 1 u") pour que Claude ait le contexte complet.
    """
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    result = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        result.append(f"\n=== Feuille: {sheet_name} ===")

        current_section = ""

        for row in ws.iter_rows(values_only=True):
            # Ignorer les lignes vides
            cells = [str(c).strip() if c is not None else '' for c in row]
            if all(c == '' for c in cells):
                continue

            first_cell = cells[0] if cells else ''
            other_cells = [c for c in cells[1:] if c and c != '0' and c != 'None']

            # Détection ligne-section : première colonne substantielle, peu d'autres données,
            # et pas de valeur numérique isolée dans les premières colonnes
            def is_numeric(s):
                try:
                    float(s.replace(',', '.').replace(' ', ''))
                    return True
                except ValueError:
                    return False

            has_price = any(is_numeric(c) for c in cells[2:5] if c)
            qte_value = cells[1] if len(cells) > 1 else ''
            has_quantity = is_numeric(qte_value)
            is_section = (
                len(first_cell) > 8
                and not has_quantity
                and len(other_cells) < 3
                and not has_price
            )

            if is_section:
                current_section = first_cell
                result.append(f"\n[SECTION] {first_cell}")
            else:
                content = ' | '.join(c for c in cells if c and c != 'None')
                if content:
                    if current_section and first_cell and first_cell != current_section:
                        result.append(f"{current_section} > {content}")
                    elif content:
                        result.append(content)

    return '\n'.join(result)


def parse_docx(file_bytes):
    """
    Parse Word document avec chunking par styles Heading (Titre 1/2/3).
    Reconstruit la structure documentaire pour faciliter l'extraction de sections.
    """
    doc = Document(io.BytesIO(file_bytes))
    result = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower()

        # Ignorer la table des matières
        if any(s in style_name for s in ['toc ', 'toc\t', 'table des mat']):
            continue

        if any(s in style_name for s in ['heading 1', 'titre 1', 'heading1']):
            result.append(f"\n## {text}")
        elif any(s in style_name for s in ['heading 2', 'titre 2', 'heading2']):
            result.append(f"\n### {text}")
        elif any(s in style_name for s in ['heading 3', 'titre 3', 'heading3']):
            result.append(f"\n#### {text}")
        elif any(s in style_name for s in ['heading 4', 'titre 4', 'heading4']):
            result.append(f"\n##### {text}")
        else:
            result.append(text)

    # Traiter les tableaux
    for table in doc.tables:
        result.append("\n[TABLEAU]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = ' | '.join(c for c in cells if c)
            if line:
                result.append(line)

    return '\n'.join(result)


def parse_pdf(file_bytes):
    """
    Parse PDF avec pdfplumber pour meilleure extraction tabulaire.
    Fallback sur extraction texte brut si pas de tableaux détectés.
    """
    result = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                result.append(f"\n--- Page {i + 1} ---")
                for table in tables:
                    for row in table:
                        cells = [str(c).strip() if c else '' for c in row]
                        non_empty = [c for c in cells if c]
                        if non_empty:
                            result.append(' | '.join(non_empty))
            else:
                text = page.extract_text()
                if text:
                    result.append(f"\n--- Page {i + 1} ---")
                    result.append(text)

    return '\n'.join(result)


@app.route('/parse/xlsx', methods=['POST'])
def route_xlsx():
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'error': 'content (base64) requis'}), 400
        file_bytes = base64.b64decode(data['content'])
        texte = parse_xlsx(file_bytes)
        return jsonify({'texte': texte})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/parse/docx', methods=['POST'])
def route_docx():
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'error': 'content (base64) requis'}), 400
        file_bytes = base64.b64decode(data['content'])
        texte = parse_docx(file_bytes)
        return jsonify({'texte': texte})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/parse/pdf', methods=['POST'])
def route_pdf():
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'error': 'content (base64) requis'}), 400
        file_bytes = base64.b64decode(data['content'])
        texte = parse_pdf(file_bytes)
        return jsonify({'texte': texte})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})


def _criticite_moeai_to_synthek(code):
    """Convertit le code MOE.AI en criticité Synthek — V2.1 complet."""
    return {
        'C01': 'MAJEUR',
        'C02': 'MINEUR',
        'C03': 'CRITIQUE',
        'C04': 'MAJEUR',
        'C05': 'CRITIQUE',
        'INCERTAIN': 'INCERTAIN',
    }.get(code, 'MINEUR')


@app.route('/compare/cctp-dpgf', methods=['POST'])
def route_compare_cctp_dpgf():
    """
    Compare un CCTP et un DPGF selon les règles MOE.AI.

    Body JSON :
    {
        "cctp": "<base64 du .docx>",
        "dpgf": "<base64 du .xlsx>",
        "config": {
            "projet": "Mon projet",
            "mapping_batiments": {
                "CCTP_section_3": ["BAT A", "BAT B"]
            },
            "programme": [
                {
                    "nom": "Bâtiment A",
                    "section_cctp": "CCTP_section_3",
                    "feuilles_dpgf": ["BAT A"],
                    "nb_logements_total": 12,
                    "types_logements": {"Accession": 8, "BRS": 4},
                    "systeme_chauffage": "Chaudière gaz N0/N1 + PAC N2"
                }
            ]
        }
    }

    Retourne :
    {
        "alertes": [ ... ],
        "nb_alertes": 5,
        "nb_conformes": 245
    }
    """
    try:
        data = request.get_json()
        if not data or 'cctp' not in data or 'dpgf' not in data:
            return jsonify({'error': 'cctp (base64) et dpgf (base64) requis'}), 400

        cctp_bytes = base64.b64decode(data['cctp'])
        dpgf_bytes = base64.b64decode(data['dpgf'])
        config = data.get('config', {})

        if not config.get('mapping_batiments'):
            config['mapping_batiments'] = {}

        # Extraire le programme bâtiment si présent
        programme = extraire_programme(config) if 'programme' in config else None

        articles = extraire_cctp(cctp_bytes, config)
        lignes = extraire_dpgf(dpgf_bytes, config)
        alertes = detecter_alertes(
            articles, lignes, config,
            utiliser_ia=False,
            programme=programme,
        )

        return jsonify({
            'alertes': [
                {
                    'code': a.code,
                    'criticite': _criticite_moeai_to_synthek(a.code),
                    'confiance': a.confiance,
                    'batiment': a.batiment,
                    'cctp_section': a.cctp_section,
                    'cctp_texte': a.cctp_texte,
                    'dpgf_feuille': a.dpgf_feuille,
                    'dpgf_ligne': a.dpgf_ligne,
                    'dpgf_texte': a.dpgf_texte,
                    'motif': a.motif,
                    'regle': a.regle,
                    'methode': a.methode,
                }
                for a in alertes
            ],
            'nb_alertes': len(alertes),
            'nb_conformes': max(0, len(articles) - len(alertes)),
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5001, debug=False)
